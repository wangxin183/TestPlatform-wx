"""Stage 1: Document ingestion — extracts text and converts to unified Markdown.

All formats (docx, pdf, xlsx, txt, md, json, yaml) are converted
to Markdown so downstream stages work with a single format.
"""

from __future__ import annotations

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from src.core.models.models import Document
from src.pipeline.stages.base import AbstractStage, StageInput, StageOutput
from src.utils.file_storage import exists, read
from src.utils.logging_config import get_logger
from src.utils.stage_logger import get_stage_logger

logger = get_logger(__name__)

GARBLED_THRESHOLD = 0.15
NONPRINTABLE_THRESHOLD = 0.25
BINARY_FORMATS = {"docx", "pdf", "xlsx"}

BINARY_SIGNATURES = [
    b"PK\x03\x04",       # ZIP / docx / xlsx
    b"%PDF",             # PDF
    b"\xd0\xcf\x11\xe0", # OLE2 / old doc/xls
    b"\x89PNG", b"\xff\xd8\xff", b"GIF8",
    b"\x00\x00\x01\x00", b"RIFF", b"\x1f\x8b",
    b"BZh", b"\x00asm",
]

# Max chars per doc stored in DB (full text stays in context)
DB_TEXT_LIMIT = 200_000


class IngestionStage(AbstractStage):
    stage_name = "ingestion"

    @classmethod
    def required_context_fields(cls) -> list[str]:
        return ["document_ids"]

    @classmethod
    def produced_context_fields(cls) -> list[str]:
        return ["raw_texts"]

    def __init__(self, db_session: AsyncSession):
        self._db = db_session

    async def execute(self, stage_input: StageInput) -> StageOutput:
        context = stage_input.context
        pid = stage_input.pipeline_id
        slog = get_stage_logger(pid, self.stage_name)
        slog.info(f"========== 文档摄入阶段开始 ==========")
        slog.info(f"文档数量: {len(context.document_ids)}")
        slog.info(f"文档ID列表: {[did[:8] for did in context.document_ids]}")
        
        raw_texts: dict[str, str] = {}
        documents: list[dict] = []
        errors: list[str] = []

        for doc_id in context.document_ids:
            doc_info = {"doc_id": doc_id, "status": "unknown"}

            result = await self._db.execute(
                select(Document).where(Document.id == doc_id)
            )
            doc = result.scalar_one_or_none()
            if doc is None:
                errors.append(f"文档 {doc_id[:8]}... 在数据库中不存在")
                doc_info["status"] = "error"
                doc_info["error"] = "文档记录未找到"
                documents.append(doc_info)
                continue

            doc_info["filename"] = doc.filename
            doc_info["file_type"] = doc.file_type

            if not await exists(doc.file_path):
                errors.append(f"文件不存在: {doc.file_path}")
                doc_info["status"] = "error"
                doc_info["error"] = f"文件未找到: {doc.file_path}"
                documents.append(doc_info)
                continue

            content = await read(doc.file_path)
            if content is None:
                errors.append(f"文件内容为空: {doc.filename}")
                doc_info["status"] = "error"
                doc_info["error"] = "文件内容为空"
                documents.append(doc_info)
                continue

            # Detect unhandled binary for text-type files
            if doc.file_type not in BINARY_FORMATS and self._has_binary_signature(content):
                errors.append(f"文件为二进制格式，无法作为文本读取: {doc.filename}")
                doc_info["status"] = "error"
                doc_info["error"] = "文件为二进制格式，不支持解析"
                documents.append(doc_info)
                continue

            # ── Convert to Markdown ──
            md_text = await self._to_markdown(content, doc.filename, doc.file_type)
            if md_text is None:
                errors.append(f"文件内容解析失败: {doc.filename}")
                doc_info["status"] = "error"
                doc_info["error"] = "文件内容解析失败，无法提取有效文本"
                documents.append(doc_info)
                continue

            # Garbled check for text-type files
            if doc.file_type not in BINARY_FORMATS and self._is_garbled(md_text):
                errors.append(f"文件内容乱码: {doc.filename}")
                doc_info["status"] = "error"
                doc_info["error"] = "文件内容乱码，无法正确解析文档编码"
                doc_info["content_preview"] = md_text[:500]
                documents.append(doc_info)
                continue

            # Success
            raw_texts[doc_id] = md_text
            doc_info["status"] = "success"
            doc_info["content_length"] = len(md_text)
            doc_info["content_preview"] = md_text[:3000]

            doc.status = "parsed"
            doc.raw_text = md_text[:DB_TEXT_LIMIT]
            self._db.add(doc)
            documents.append(doc_info)

        await self._db.commit()
        context.raw_texts = raw_texts

        success_count = sum(1 for d in documents if d["status"] == "success")
        error_count = len(errors)
        
        slog.info(f"文档处理完成: 成功={success_count}, 失败={error_count}")
        for d in documents:
            slog.info(f"  文档 {d.get('filename', d['doc_id'][:8])}: {d['status']}" + 
                      (f" (内容长度={d.get('content_length', 0)})" if d['status'] == 'success' else 
                       f" (错误={d.get('error', '')})" if d.get('error') else ""))
        
        if errors:
            logger.warning(
                "ingestion_errors",
                pipeline_id=stage_input.pipeline_id,
                success=success_count,
                errors=error_count,
                error_details=errors,
            )

        if not raw_texts:
            return StageOutput(
                stage_name=self.stage_name,
                status="failed",
                data={
                    "document_count": len(context.document_ids),
                    "success_count": 0,
                    "error_count": error_count,
                    "documents": documents,
                },
                error="; ".join(errors) if errors else "没有文档被成功导入",
            )

        return StageOutput(
            stage_name=self.stage_name,
            status="completed",
            data={
                "document_count": len(context.document_ids),
                "success_count": success_count,
                "error_count": error_count,
                "documents": documents,
            },
        )

    # ── Format converters ──

    async def _to_markdown(self, content: bytes, filename: str, file_type: str) -> str | None:
        if file_type == "docx":
            return self._docx_to_md(content)
        elif file_type == "pdf":
            return self._pdf_to_md(content)
        elif file_type == "xlsx":
            return self._xlsx_to_md(content)
        elif file_type in ("json", "openapi_json"):
            return self._json_to_md(content)
        elif file_type in ("yaml", "openapi_yaml"):
            return self._yaml_to_md(content)
        elif file_type == "md":
            return self._decode_text(content, filename)
        else:
            return self._text_to_md(content, filename)

    # ── docx ──

    def _docx_to_md(self, content: bytes) -> str | None:
        from io import BytesIO
        from docx import Document as DocxDocument

        try:
            doc = DocxDocument(BytesIO(content))
            parts = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    parts.append("")
                    continue

                style_name = (para.style.name if para.style else "").lower()
                is_heading = "heading" in style_name or "title" in style_name

                if is_heading:
                    level = self._guess_heading_level(para)
                    prefix = "#" * level
                    parts.append(f"{prefix} {text}")
                elif para.runs and para.runs[0].bold:
                    if len(text) < 80 and not text.endswith("。"):
                        parts.append(f"### {text}")
                    else:
                        parts.append(text)
                else:
                    parts.append(text)

            # Tables → Markdown tables
            for table in doc.tables:
                parts.append("")
                rows = []
                for row in table.rows:
                    cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                    rows.append(cells)
                if rows:
                    parts.append(self._rows_to_md_table(rows))

            return "\n\n".join(parts) or None
        except Exception as exc:
            logger.warning("docx_to_md_failed", error=str(exc))
            return None

    def _guess_heading_level(self, para) -> int:
        """Map Word heading style to markdown heading level."""
        style_name = (para.style.name if para.style else "").lower()
        if "heading 1" in style_name or "title" in style_name:
            return 1
        elif "heading 2" in style_name:
            return 2
        elif "heading 3" in style_name:
            return 3
        # Font size heuristic
        if para.runs:
            run = para.runs[0]
            if run.font.size:
                pt = run.font.size.pt
                if pt >= 18:
                    return 1
                elif pt >= 14:
                    return 2
                elif pt >= 12:
                    return 3
        return 2

    # ── PDF ──

    def _pdf_to_md(self, content: bytes) -> str | None:
        from io import BytesIO
        from PyPDF2 import PdfReader

        try:
            reader = PdfReader(BytesIO(content))
            parts = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    parts.append(f"<!-- page {i+1} -->\n\n{text.strip()}")
            return "\n\n".join(parts) or None
        except Exception as exc:
            logger.warning("pdf_to_md_failed", error=str(exc))
            return None

    # ── xlsx ──

    def _xlsx_to_md(self, content: bytes) -> str | None:
        from io import BytesIO
        from openpyxl import load_workbook

        try:
            wb = load_workbook(BytesIO(content), read_only=True, data_only=True)
            parts = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows = list(ws.iter_rows(values_only=True))
                if not rows:
                    continue
                parts.append(f"## {sheet_name}")
                parts.append(self._rows_to_md_table(rows))
            wb.close()
            return "\n\n".join(parts) or None
        except Exception as exc:
            logger.warning("xlsx_to_md_failed", error=str(exc))
            return None

    # ── JSON / YAML ──

    def _json_to_md(self, content: bytes) -> str | None:
        try:
            import json as _json
            data = _json.loads(content)
            text = _json.dumps(data, ensure_ascii=False, indent=2)
            if len(text) < 50:
                return None
            return text
        except Exception:
            return self._decode_text(content, "<json>")

    def _yaml_to_md(self, content: bytes) -> str | None:
        text = self._decode_text(content, "<yaml>")
        if text:
            return text
        return None

    # ── Plain text ──

    def _text_to_md(self, content: bytes, filename: str) -> str | None:
        text = self._decode_text(content, filename)
        if text is None:
            return None
        # If text already has markdown headings, keep as-is
        if any(line.strip().startswith("#") for line in text.split("\n")):
            return text
        return text

    # ── Helpers ──

    def _rows_to_md_table(self, rows: list[tuple]) -> str:
        """Convert rows (list of tuples) to a Markdown table string."""
        if not rows:
            return ""
        max_cols = max(len(r) for r in rows)

        def norm(r):
            padded = list(r) + [""] * (max_cols - len(r))
            return [str(c) if c is not None else "" for c in padded]

        header = norm(rows[0])
        lines = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join(["---"] * max_cols) + " |",
        ]
        for r in rows[1:]:
            cells = norm(r)
            if all(c == "" for c in cells):
                continue
            lines.append("| " + " | ".join(cells) + " |")
        return "\n".join(lines)

    def _decode_text(self, content: bytes, filename: str) -> str | None:
        if content.startswith(b"\xef\xbb\xbf"):
            content = content[3:]
        try:
            text = content.decode("utf-8")
            if "�" not in text:
                return text
        except UnicodeDecodeError:
            pass
        for bom, enc in [(b"\xff\xfe", "utf-16-le"), (b"\xfe\xff", "utf-16-be")]:
            if content.startswith(bom):
                try:
                    return content[2:].decode(enc)
                except UnicodeDecodeError:
                    pass
        for encoding in ["gb18030", "gbk", "gb2312"]:
            try:
                return content.decode(encoding)
            except (UnicodeDecodeError, LookupError):
                continue
        try:
            return content.decode("utf-8", errors="replace")
        except Exception:
            return None

    def _has_binary_signature(self, content: bytes) -> bool:
        for sig in BINARY_SIGNATURES:
            if content.startswith(sig):
                return True
        return False

    def _is_garbled(self, text: str) -> bool:
        if not text or len(text.strip()) < 10:
            return True
        if text.count("�") / max(len(text), 1) > GARBLED_THRESHOLD:
            return True
        non_printable = sum(1 for c in text if ord(c) < 32 and c not in "\n\r\t")
        if non_printable / max(len(text), 1) > NONPRINTABLE_THRESHOLD:
            return True
        return False
