"""文档格式转换器 — 将各种文档格式转换为 Markdown。

从 IngestionStage 中提取出来的纯函数版本，无 DB 依赖。
用于需求分析节点的文档摄取。

Supported formats: docx, pdf, xlsx, json, yaml, md, txt

Usage:
    from src.utils.document_converter import convert_to_markdown

    md_text = convert_to_markdown(content_bytes, filename="需求文档.docx", file_type="docx")
"""

from __future__ import annotations

import json as _json
from pathlib import Path

from src.utils.logging_config import get_logger

logger = get_logger(__name__)

# ============================================================
# 公共接口
# ============================================================


def convert_to_markdown(
    content: bytes,
    filename: str = "",
    file_type: str = "",
) -> str | None:
    """将文档内容转换为 Markdown 文本。

    Args:
        content: 文档的原始字节内容
        filename: 原始文件名（用于编码检测）
        file_type: 文件类型（docx/pdf/xlsx/json/yaml/md/txt）

    Returns:
        Markdown 文本；如果转换失败则返回 None
    """
    if not file_type and filename:
        file_type = _guess_type(filename)

    if file_type == "docx":
        return _docx_to_md(content)
    elif file_type == "pdf":
        return _pdf_to_md(content)
    elif file_type == "xlsx":
        return _xlsx_to_md(content)
    elif file_type in ("json", "openapi_json"):
        return _json_to_md(content)
    elif file_type in ("yaml", "openapi_yaml"):
        return _yaml_to_md(content)
    elif file_type == "md":
        return _decode_text(content, filename)
    else:
        return _text_to_md(content, filename)


def detect_file_type(filename: str | Path) -> str:
    """根据文件扩展名确定文件类型。

    Returns:
        docx / pdf / xlsx / json / yaml / md / txt
    """
    ext = Path(filename).suffix.lower()
    type_map = {
        ".docx": "docx",
        ".doc": "docx",
        ".pdf": "pdf",
        ".xlsx": "xlsx",
        ".xls": "xlsx",
        ".json": "openapi_json",
        ".yaml": "openapi_yaml",
        ".yml": "openapi_yaml",
        ".md": "md",
        ".txt": "txt",
    }
    return type_map.get(ext, "txt")


def has_binary_signature(content: bytes) -> bool:
    """检测内容是否为无法处理的二进制格式。"""
    signatures = [
        b"PK\x03\x04",       # ZIP / docx / xlsx
        b"%PDF",             # PDF
        b"\xd0\xcf\x11\xe0", # OLE2
        b"\x89PNG",
        b"\xff\xd8\xff",
        b"GIF8",
        b"\x00\x00\x01\x00",
        b"RIFF",
        b"\x1f\x8b",
        b"BZh",
        b"\x00asm",
    ]
    return any(content.startswith(sig) for sig in signatures)


# ============================================================
# 内部转换函数
# ============================================================


def _docx_to_md(content: bytes) -> str | None:
    """将 .docx 文档转换为 Markdown。"""
    from io import BytesIO
    from docx import Document as DocxDocument

    try:
        doc = DocxDocument(BytesIO(content))
        parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                parts.append("")
                continue

            style_name = (para.style.name if para.style else "").lower()
            is_heading = "heading" in style_name or "title" in style_name

            if is_heading:
                level = _guess_heading_level(para)
                parts.append(f"{'#' * level} {text}")
            elif para.runs and para.runs[0].bold:
                if len(text) < 80 and not text.endswith("。"):
                    parts.append(f"### {text}")
                else:
                    parts.append(text)
            else:
                parts.append(text)

        # 表格 → Markdown 表格
        for table in doc.tables:
            parts.append("")
            rows = [
                [cell.text.replace("\n", " ").strip() for cell in row.cells]
                for row in table.rows
            ]
            if rows:
                parts.append(_rows_to_md_table(rows))

        return "\n\n".join(parts) or None
    except Exception as exc:
        logger.warning("docx_to_md_failed", error=str(exc))
        return None


def _pdf_to_md(content: bytes) -> str | None:
    """将 PDF 文档转换为 Markdown。"""
    from io import BytesIO
    from PyPDF2 import PdfReader

    try:
        reader = PdfReader(BytesIO(content))
        parts = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                parts.append(f"<!-- page {i+1} -->\n\n{text.strip()}")
        return "\n\n".join(parts) or None
    except Exception as exc:
        logger.warning("pdf_to_md_failed", error=str(exc))
        return None


def _xlsx_to_md(content: bytes) -> str | None:
    """将 Excel 文档转换为 Markdown 表格。"""
    from io import BytesIO
    from openpyxl import load_workbook

    try:
        wb = load_workbook(BytesIO(content), read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue
            parts.append(f"## {sheet_name}")
            parts.append(_rows_to_md_table(rows))
        wb.close()
        return "\n\n".join(parts) or None
    except Exception as exc:
        logger.warning("xlsx_to_md_failed", error=str(exc))
        return None


def _json_to_md(content: bytes) -> str | None:
    """将 JSON 内容转换为格式化文本。"""
    try:
        data = _json.loads(content)
        text = _json.dumps(data, ensure_ascii=False, indent=2)
        if len(text) < 50:
            return None
        return text
    except Exception:
        return _decode_text(content, "<json>")


def _yaml_to_md(content: bytes) -> str | None:
    """将 YAML 内容转换为文本。"""
    text = _decode_text(content, "<yaml>")
    return text if text else None


def _text_to_md(content: bytes, filename: str) -> str | None:
    """将纯文本转换为 Markdown（编码检测后原样返回）。"""
    text = _decode_text(content, filename)
    if text is None:
        return None
    return text


# ============================================================
# 辅助函数
# ============================================================


def _rows_to_md_table(rows: list[tuple]) -> str:
    """将行列表转换为 Markdown 表格字符串。"""
    if not rows:
        return ""
    max_cols = max(len(r) for r in rows)

    def norm(r):
        padded = list(r) + [""] * (max_cols - len(r))
        return [str(c) if c is not None else "" for c in padded]

    header = norm(rows[0])
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * max_cols) + " |",
    ]
    for r in rows[1:]:
        cells = norm(r)
        if all(c == "" for c in cells):
            continue
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def _guess_heading_level(para) -> int:
    """根据 Word 样式推断 Markdown 标题级别。"""
    style_name = (para.style.name if para.style else "").lower()
    if "heading 1" in style_name or "title" in style_name:
        return 1
    elif "heading 2" in style_name:
        return 2
    elif "heading 3" in style_name:
        return 3
    # 字号启发式检测
    if para.runs:
        run = para.runs[0]
        if run.font.size:
            pt = run.font.size.pt
            if pt >= 18:
                return 1
            elif pt >= 14:
                return 2
            elif pt >= 12:
                return 3
    return 2


def _decode_text(content: bytes, filename: str = "") -> str | None:
    """尝试多种编码解码文本内容。"""
    # 去掉 UTF-8 BOM
    if content.startswith(b"\xef\xbb\xbf"):
        content = content[3:]

    # UTF-8
    try:
        text = content.decode("utf-8")
        if "�" not in text:
            return text
    except UnicodeDecodeError:
        pass

    # UTF-16（带 BOM）
    for bom, enc in [(b"\xff\xfe", "utf-16-le"), (b"\xfe\xff", "utf-16-be")]:
        if content.startswith(bom):
            try:
                return content[2:].decode(enc)
            except UnicodeDecodeError:
                pass

    # GB 编码（中文 Windows 常用）
    for encoding in ["gb18030", "gbk", "gb2312"]:
        try:
            return content.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue

    # 最后回退：UTF-8 替代字符
    try:
        return content.decode("utf-8", errors="replace")
    except Exception:
        return None


def _guess_type(filename: str) -> str:
    """根据文件名推断文件类型。"""
    ext = Path(filename).suffix.lower()
    return {
        ".docx": "docx",
        ".pdf": "pdf",
        ".xlsx": "xlsx",
        ".json": "json",
        ".yaml": "yaml",
        ".yml": "yaml",
        ".md": "md",
        ".txt": "txt",
    }.get(ext, "txt")
